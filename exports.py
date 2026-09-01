# -*- coding: utf-8 -*-
"""
EduTest — Excel (.xlsx) болон Word (.docx) тайлан үүсгэх модуль.

Энэ файл нь ЗӨВХӨН форматлалт хийнэ. Тааруулалт, бодолт, нэгтгэлийг
`domain.py` аль хэдийн хийсэн байх ёстой — энд ямар ч бизнес логик
давхардуулж бичихгүй (нэг үнэний эх сурвалж).

Гадаад хамаарал:
    openpyxl     — Excel
    python-docx  — Word

Файлын нэр:
    PHR201_Group-A_PrePost_2026-08-31.xlsx
    PHR201_Group-A_PrePost_Report_2026-08-31.docx
"""

from __future__ import annotations

import io
import re
import unicodedata
from datetime import date

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------
# График (заавал биш хамаарал)
# ---------------------------------------------------------------------
# matplotlib байхгүй бол тайлан ГРАФИКГҮЙГЭЭР үүсч, бусад бүх зүйл хэвээр
# ажиллана. Ингэснээр график зурах чадвар нь Word экспортын бүхэл бүтэн
# ажиллагааг зогсоохгүй.
try:
    import matplotlib
    matplotlib.use("Agg")          # дэлгэцгүй сервер дээр ажиллана
    import matplotlib.pyplot as plt
    HAS_CHARTS = True
except ImportError:                # pragma: no cover
    plt = None
    HAS_CHARTS = False

# Вэб дээрх дизайн токентой ижил өнгө — тайлан ба дэлгэц ижил харагдана.
C_PRIMARY = "#4F46E5"
C_SUCCESS = "#0F766E"
C_DANGER = "#B42318"
C_MUTED = "#6B7793"
C_GRID = "#E3E7F0"
C_TEXT = "#16203A"

# ---------------------------------------------------------------------
# Нийтлэг тохиргоо
# ---------------------------------------------------------------------
STATUS_MN = {"matched": "Тааруулсан", "pre_only": "Зөвхөн оролт", "post_only": "Зөвхөн гаралт"}

_HEADER_FILL = PatternFill("solid", fgColor="24365A")
_HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
_TITLE_FONT = Font(bold=True, size=13, color="16203A")
_LABEL_FONT = Font(bold=True, size=10, color="414E6C")
_THIN = Side(style="thin", color="D3D9E6")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)

_UP_FILL = PatternFill("solid", fgColor="E6F4F2")
_DOWN_FILL = PatternFill("solid", fgColor="FDECEA")
_FLAT_FILL = PatternFill("solid", fgColor="F1F3F9")


def safe_filename_part(value: str | None, fallback: str = "EduTest") -> str:
    """
    Файлын нэрэнд тавихад аюулгүй болгоно.

    Кирилл үсгийг латин руу хөрвүүлэхгүй — оронд нь ASCII бус тэмдэгтийг
    зайгаар сольж, үлдсэнийг зурааслана. Ингэснээр Windows/macOS/Linux
    бүгд дээр асуудалгүй нээгдэнэ.
    """
    text = unicodedata.normalize("NFKD", str(value or "")).encode("ascii", "ignore").decode()
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-")
    return text or fallback


def build_filename(course_code, group_name, suffix: str, extension: str,
                   on: date | None = None) -> str:
    """Жишээ: PHR201_Group-A_PrePost_2026-08-31.xlsx"""
    stamp = (on or date.today()).isoformat()
    parts = [
        safe_filename_part(course_code, "COURSE"),
        safe_filename_part(group_name, "All-Groups"),
        suffix,
        stamp,
    ]
    return "_".join(parts) + "." + extension


# =====================================================================
# EXCEL
# =====================================================================
def _write_header(ws, headers: list, row: int = 1) -> None:
    for col, title in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col, value=title)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = _BORDER


def _autofit(ws, min_width: int = 9, max_width: int = 42) -> None:
    """Баганын өргөнийг агуулгад тохируулна (openpyxl-д автомат байхгүй)."""
    for column_cells in ws.columns:
        longest = 0
        letter = get_column_letter(column_cells[0].column)
        for cell in column_cells:
            if cell.value is not None:
                longest = max(longest, max(len(line) for line in str(cell.value).split("\n")))
        ws.column_dimensions[letter].width = min(max(longest + 3, min_width), max_width)


def build_pair_workbook(pair: dict, rows: list, summary: dict,
                        tests: dict | None = None) -> bytes:
    """
    Pre/Post харьцуулалтыг гурван хуудастай Excel болгоно:
        Summary            — нэгтгэсэн үзүүлэлт
        Student Results    — тааруулсан оюутан бүрийн мөр
        Unmatched Attempts — зөвхөн Оролт эсвэл зөвхөн Гаралт өгсөн оролдлого
    """
    wb = Workbook()

    # ---------------- Summary ----------------
    ws = wb.active
    ws.title = "Summary"
    ws["A1"] = "EduTest — Оролт/Гаралтын тестийн нэгтгэл"
    ws["A1"].font = _TITLE_FONT
    ws.merge_cells("A1:B1")

    facts = [
        ("Хичээл", f"{pair.get('course_name', '')} ({pair.get('course_code', '')})"),
        ("Тестийн хос", pair.get("name", "")),
        ("Тайлангийн огноо", date.today().isoformat()),
        ("", ""),
        ("Нийт мөр", summary.get("total_rows", 0)),
        ("Тааруулсан оюутан", summary.get("matched_count", 0)),
        ("Зөвхөн оролт өгсөн", summary.get("pre_only_count", 0)),
        ("Зөвхөн гаралт өгсөн", summary.get("post_only_count", 0)),
        ("", ""),
        ("Оролтын дундаж, %", summary.get("avg_pre")),
        ("Гаралтын дундаж, %", summary.get("avg_post")),
        ("Дундаж ахиц, %", summary.get("avg_delta")),
        ("", ""),
        ("Ахисан оюутан", summary.get("improved_count", 0)),
        ("Өөрчлөлтгүй", summary.get("same_count", 0)),
        ("Буурсан", summary.get("declined_count", 0)),
        ("Нэрийн зөрчилтэй мөр", summary.get("conflict_count", 0)),
    ]
    row = 3
    for label, value in facts:
        if label:
            ws.cell(row=row, column=1, value=label).font = _LABEL_FONT
            cell = ws.cell(row=row, column=2, value=value if value is not None else "—")
            if isinstance(value, (int, float)):
                cell.number_format = "0.0" if isinstance(value, float) else "0"
        row += 1

    matched = summary.get("matched_count", 0) or 0
    improved = summary.get("improved_count", 0) or 0
    ws.cell(row=row + 1, column=1, value="Ахисан оюутны хувь").font = _LABEL_FONT
    pct_cell = ws.cell(row=row + 1, column=2,
                       value=round(improved * 100 / matched, 1) if matched else 0)
    pct_cell.number_format = "0.0"

    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 34

    # ---------------- Student Results ----------------
    ws2 = wb.create_sheet("Student Results")
    headers = ["№", "Оюутны код", "Овог нэр", "Бүлэг", "Имэйл (заавал бус)",
               "Оролт, %", "Гаралт, %", "Өөрчлөлт", "Өсөлтийн хувь, %",
               "Тааралт", "Нэрийн зөрчил"]
    _write_header(ws2, headers)

    matched_rows = [r for r in rows if r["status"] == "matched"]
    for index, r in enumerate(matched_rows, start=1):
        pre = r["pre_percent"]
        post = r["post_percent"]
        delta = r["delta_percent"]
        growth = round(delta * 100 / pre, 1) if (pre and delta is not None) else None
        values = [
            index, r["student_code"], r["full_name"], r["class_group_name"],
            r["email"] or "", pre, post, delta, growth,
            STATUS_MN.get(r["status"], r["status"]),
            " / ".join(r["conflicting_names"]) if r["name_conflict"] else "",
        ]
        for col, value in enumerate(values, start=1):
            cell = ws2.cell(row=index + 1, column=col, value=value)
            cell.border = _BORDER
            if col in (6, 7, 8, 9):
                cell.number_format = "0.0" if col == 9 else "0"
                cell.alignment = Alignment(horizontal="right")
        # Ахицыг өнгөөр ялгана: эерэг ногоон, сөрөг улаан, тэг саарал.
        if delta is not None:
            fill = _UP_FILL if delta > 0 else (_DOWN_FILL if delta < 0 else _FLAT_FILL)
            ws2.cell(row=index + 1, column=8).fill = fill

    ws2.freeze_panes = "A2"
    if matched_rows:
        ws2.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(matched_rows) + 1}"
    _autofit(ws2)

    # ---------------- Unmatched Attempts ----------------
    ws3 = wb.create_sheet("Unmatched Attempts")
    headers3 = ["№", "Оюутны код", "Овог нэр", "Бүлэг", "Имэйл (заавал бус)",
                "Оролт, %", "Гаралт, %", "Төлөв", "Тайлбар"]
    _write_header(ws3, headers3)

    unmatched = [r for r in rows if r["status"] != "matched"]
    for index, r in enumerate(unmatched, start=1):
        note = ("Гаралтын тест өгөөгүй" if r["status"] == "pre_only"
                else "Оролтын тест өгөөгүй")
        values = [
            index, r["student_code"], r["full_name"], r["class_group_name"],
            r["email"] or "", r["pre_percent"], r["post_percent"],
            STATUS_MN.get(r["status"], r["status"]), note,
        ]
        for col, value in enumerate(values, start=1):
            cell = ws3.cell(row=index + 1, column=col, value=value)
            cell.border = _BORDER
            if col in (6, 7):
                cell.number_format = "0"
                cell.alignment = Alignment(horizontal="right")

    ws3.freeze_panes = "A2"
    if unmatched:
        ws3.auto_filter.ref = f"A1:{get_column_letter(len(headers3))}{len(unmatched) + 1}"
    _autofit(ws3)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# =====================================================================
# ГРАФИК — Word тайланд шигтгэх PNG зургууд
# ---------------------------------------------------------------------
# Эдгээр нь Шинжилгээ хуудасны Chart.js график ЯГ ижил өгөгдөл, ижил
# өнгөөр зурагдана. Тооцоолол давхардуулаагүй — `rows` болон `summary`
# нь domain.py-аас ирсэн бэлэн үр дүн.
# =====================================================================
def _fmt_num(value) -> str:
    """64.0 -> '64',  64.5 -> '64.5' — илүү тэг харуулахгүй."""
    return str(int(value)) if float(value) == int(value) else str(value)


def _style_axes(ax, *, ylabel=None, percent=False, integer_ticks=False):
    """Бүх графикт нэг ижил цэвэрхэн загвар өгнө."""
    ax.set_facecolor("white")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(C_GRID)
    ax.spines["bottom"].set_color(C_GRID)
    ax.tick_params(colors=C_MUTED, labelsize=9)
    ax.yaxis.grid(True, color=C_GRID, linewidth=0.8)
    ax.set_axisbelow(True)
    if ylabel:
        ax.set_ylabel(ylabel, color=C_MUTED, fontsize=9)
    if percent:
        ax.set_ylim(0, 100)
        ax.yaxis.set_major_formatter(lambda v, _pos: f"{int(v)}%")
    if integer_ticks:
        # Оюутны тоо бүхэл байна — 0.5, 1.5 гэсэн заалт утгагүй.
        from matplotlib.ticker import MaxNLocator
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))


def _figure_png(fig) -> io.BytesIO:
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def chart_average(summary: dict):
    """Оролт ба Гаралтын дундаж — хоёр багана."""
    pre = summary.get("avg_pre") or 0
    post = summary.get("avg_post") or 0
    fig, ax = plt.subplots(figsize=(5.2, 3.0))
    bars = ax.bar(["Оролт", "Гаралт"], [pre, post],
                  color=[C_PRIMARY, C_SUCCESS], width=0.5)
    for bar, value in zip(bars, (pre, post)):
        ax.annotate(f"{_fmt_num(value)}%", (bar.get_x() + bar.get_width() / 2, value),
                    ha="center", va="bottom", fontsize=10, fontweight="bold",
                    color=C_TEXT, xytext=(0, 3), textcoords="offset points")
    _style_axes(ax, percent=True)
    ax.set_title("Оролт ба Гаралтын дундаж", fontsize=11, color=C_TEXT, pad=12)
    return _figure_png(fig)


def chart_improvement(summary: dict):
    """Ахисан / Өөрчлөлтгүй / Буурсан оюутны тоо."""
    values = [summary.get("improved_count", 0), summary.get("same_count", 0),
              summary.get("declined_count", 0)]
    fig, ax = plt.subplots(figsize=(5.2, 3.0))
    bars = ax.bar(["Ахисан", "Өөрчлөлтгүй", "Буурсан"], values,
                  color=[C_SUCCESS, C_MUTED, C_DANGER], width=0.55)
    for bar, value in zip(bars, values):
        ax.annotate(str(value), (bar.get_x() + bar.get_width() / 2, value),
                    ha="center", va="bottom", fontsize=10, fontweight="bold",
                    color=C_TEXT, xytext=(0, 3), textcoords="offset points")
    _style_axes(ax, ylabel="Оюутны тоо", integer_ticks=True)
    ax.set_ylim(0, max(values + [1]) * 1.25)
    ax.set_title("Ахицын ангилал", fontsize=11, color=C_TEXT, pad=12)
    return _figure_png(fig)


def chart_students(rows: list, limit: int = 30):
    """
    Оюутан тус бүрийн Оролт/Гаралт.
    Хэт олон оюутантай үед графикийг уншигдахуйц байлгахын тулд эхний
    `limit` оюутныг л зурна (бүрэн жагсаалт хүснэгтэд байна).
    """
    matched = [r for r in rows if r["status"] == "matched"][:limit]
    if not matched:
        return None

    labels = [r["student_code"] for r in matched]
    pre = [r["pre_percent"] for r in matched]
    post = [r["post_percent"] for r in matched]
    positions = range(len(matched))
    width = 0.4

    fig, ax = plt.subplots(figsize=(max(6.4, len(matched) * 0.42), 3.4))
    ax.bar([p - width / 2 for p in positions], pre, width,
           label="Оролт", color=C_PRIMARY)
    ax.bar([p + width / 2 for p in positions], post, width,
           label="Гаралт", color=C_SUCCESS)
    ax.set_xticks(list(positions))
    ax.set_xticklabels(labels, rotation=60, ha="right", fontsize=8)
    _style_axes(ax, percent=True)
    ax.legend(frameon=False, fontsize=9, loc="upper left",
              bbox_to_anchor=(0, 1.14), ncol=2)
    ax.set_title("Оюутан тус бүрийн Оролт/Гаралт", fontsize=11,
                 color=C_TEXT, pad=26)
    return _figure_png(fig)


def chart_distribution(rows: list):
    """Онооны бүсэд ногдох оюутны тоо (Оролт ба Гаралтыг зэрэгцүүлэн)."""
    matched = [r for r in rows if r["status"] == "matched"]
    if not matched:
        return None

    buckets = ["0-40", "41-60", "61-80", "81-100"]

    def bucket_of(value):
        if value <= 40:
            return "0-40"
        if value <= 60:
            return "41-60"
        if value <= 80:
            return "61-80"
        return "81-100"

    pre_counts = dict.fromkeys(buckets, 0)
    post_counts = dict.fromkeys(buckets, 0)
    for r in matched:
        pre_counts[bucket_of(r["pre_percent"])] += 1
        post_counts[bucket_of(r["post_percent"])] += 1

    pre = [pre_counts[b] for b in buckets]
    post = [post_counts[b] for b in buckets]
    positions = range(len(buckets))
    width = 0.38

    fig, ax = plt.subplots(figsize=(5.6, 3.0))
    ax.bar([p - width / 2 for p in positions], pre, width,
           label="Оролт", color=C_PRIMARY)
    ax.bar([p + width / 2 for p in positions], post, width,
           label="Гаралт", color=C_SUCCESS)
    ax.set_xticks(list(positions))
    ax.set_xticklabels(buckets)
    _style_axes(ax, ylabel="Оюутны тоо", integer_ticks=True)
    ax.set_ylim(0, max(pre + post + [1]) * 1.3)
    ax.legend(frameon=False, fontsize=9)
    ax.set_title("Онооны тархалт", fontsize=11, color=C_TEXT, pad=12)
    return _figure_png(fig)


# =====================================================================
# WORD
# =====================================================================
def _para(doc, text, *, size=10.5, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build_pair_report(pair: dict, rows: list, summary: dict,
                      institution: str = "EduTest", *,
                      include_charts: bool = True) -> bytes:
    """
    Pre/Post харьцуулалтын Word тайлан.

    График нь Шинжилгээ хуудасны Chart.js графиктай ИЖИЛ өгөгдөл, ижил
    өнгөөр matplotlib-аар зурагдаж PNG болж шигтгэгдэнэ.
    matplotlib суулгаагүй бол тайлан графикгүйгээр бүрэн үүснэ.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _para(doc, institution, size=11, bold=True, color="4F46E5", space_after=2)
    _para(doc, "Оролт/Гаралтын тестийн харьцуулсан тайлан",
          size=17, bold=True, color="16203A", space_after=10)

    # ------- Ерөнхий мэдээлэл -------
    meta = [
        ("Хичээл", f"{pair.get('course_name', '')} ({pair.get('course_code', '')})"),
        ("Тестийн хос", pair.get("name", "")),
        ("Бүлэг", pair.get("group_label") or "Бүх бүлэг"),
        ("Тайлангийн огноо", date.today().isoformat()),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in meta:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True

    # ------- Нэгтгэсэн үзүүлэлт -------
    doc.add_heading("Нэгтгэсэн үзүүлэлт", level=1)

    def fmt(value, suffix="%"):
        return "—" if value is None else f"{value}{suffix}"

    matched = summary.get("matched_count", 0) or 0
    improved = summary.get("improved_count", 0) or 0
    summary_rows = [
        ("Нийт оролцогч (мөр)", str(summary.get("total_rows", 0))),
        ("Тааруулсан оюутан", str(matched)),
        ("Таараагүй оролдлого",
         str((summary.get("pre_only_count", 0) or 0) + (summary.get("post_only_count", 0) or 0))),
        ("Оролтын дундаж", fmt(summary.get("avg_pre"))),
        ("Гаралтын дундаж", fmt(summary.get("avg_post"))),
        ("Дундаж ахиц", fmt(summary.get("avg_delta"))),
        ("Ахисан / Өөрчлөлтгүй / Буурсан",
         f"{improved} / {summary.get('same_count', 0)} / {summary.get('declined_count', 0)}"),
        ("Ахисан оюутны хувь",
         f"{round(improved * 100 / matched, 1)}%" if matched else "—"),
    ]
    st = doc.add_table(rows=0, cols=2)
    st.style = "Light Grid Accent 1"
    for label, value in summary_rows:
        cells = st.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True

    # ------- Дүгнэлт -------
    doc.add_heading("Ерөнхий дүгнэлт", level=1)
    avg_delta = summary.get("avg_delta")
    if matched == 0:
        verdict = ("Тааруулсан оюутан байхгүй тул ахицын дүгнэлт гаргах боломжгүй. "
                   "Оролт болон Гаралтын тестийг ижил оюутны код, ижил бүлгээр "
                   "бөглүүлсэн эсэхийг шалгана уу.")
    elif avg_delta is None:
        verdict = "Дундаж ахицыг тооцох өгөгдөл хүрэлцэхгүй байна."
    elif avg_delta > 0:
        verdict = (f"Тааруулсан {matched} оюутны дундаж оноо Оролтоос Гаралт хүртэл "
                   f"{avg_delta} нэгжээр өссөн. Нийт {improved} оюутан ахисан үзүүлэлттэй "
                   f"гарсан нь сургалтын үр дүн эерэг болохыг илэрхийлж байна.")
    elif avg_delta < 0:
        verdict = (f"Тааруулсан {matched} оюутны дундаж оноо {abs(avg_delta)} нэгжээр "
                   f"буурсан байна. Гаралтын тестийн агуулга, хугацаа болон хичээлийн "
                   f"хамрах хүрээг дахин хянахыг зөвлөж байна.")
    else:
        verdict = (f"Тааруулсан {matched} оюутны дундаж оноо өөрчлөгдөөгүй байна. "
                   f"Сургалтын аргачлалыг тодотгох шаардлагатай байж болно.")
    _para(doc, verdict, space_after=10)

    if summary.get("conflict_count"):
        _para(doc,
              f"Анхааруулга: {summary['conflict_count']} мөрд ижил бүлэг, ижил оюутны "
              f"кодоор өөр нэр бүртгэгдсэн байна. Систем эдгээрийг автоматаар "
              f"нэгтгээгүй тул багш нэрийг тулган шалгах шаардлагатай.",
              bold=True, color="B45309", space_after=10)

    # ------- График -------
    matched_count = summary.get("matched_count", 0) or 0
    if include_charts and HAS_CHARTS and matched_count:
        doc.add_heading("График", level=1)
        for builder, args, width in (
            (chart_average, (summary,), 5.4),
            (chart_improvement, (summary,), 5.4),
            (chart_distribution, (rows,), 5.8),
            (chart_students, (rows,), 6.3),
        ):
            image = builder(*args)
            if image is None:
                continue
            doc.add_picture(image, width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        shown = len([r for r in rows if r["status"] == "matched"])
        if shown > 30:
            _para(doc,
                  f"Тэмдэглэл: сүүлийн графикт эхний 30 оюутан харагдаж байна "
                  f"(нийт {shown}). Бүрэн жагсаалт доорх хүснэгтэд бий.",
                  size=8.5, color="6B7793")
    elif include_charts and not HAS_CHARTS:
        _para(doc,
              "Тэмдэглэл: график зурах сан (matplotlib) суулгаагүй тул энэ "
              "тайланд график ороогүй. `pip install matplotlib` ажиллуулаад "
              "тайланг дахин татна уу.",
              size=8.5, color="B45309")

    # ------- Хүснэгт -------
    doc.add_heading("Оюутны Оролт/Гаралтын харьцуулалт", level=1)
    headers = ["№", "Оюутны код", "Овог нэр", "Бүлэг", "Оролт", "Гаралт", "Өөрчлөлт", "Тааралт"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Medium Shading 1 Accent 1"
    for index, title in enumerate(headers):
        cell = tbl.rows[0].cells[index]
        cell.text = title
        cell.paragraphs[0].runs[0].bold = True

    for index, r in enumerate(rows, start=1):
        cells = tbl.add_row().cells
        pre = r["pre_percent"]
        post = r["post_percent"]
        delta = r["delta_percent"]
        values = [
            str(index), r["student_code"], r["full_name"], r["class_group_name"],
            "—" if pre is None else f"{pre}%",
            "—" if post is None else f"{post}%",
            "—" if delta is None else (f"+{delta}%" if delta > 0 else f"{delta}%"),
            STATUS_MN.get(r["status"], r["status"]),
        ]
        for col, value in enumerate(values):
            cells[col].text = value
        # Ахицыг өнгөөр: эерэг ногоон, сөрөг улаан, тэг саарал.
        if delta is not None:
            run = cells[6].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(
                "0F766E" if delta > 0 else ("B42318" if delta < 0 else "6B7793")
            )

    if not rows:
        _para(doc, "Харьцуулах өгөгдөл алга байна.", color="6B7793")

    _para(doc, "", space_after=12)
    _para(doc,
          "Тааруулалтын түлхүүр: ангийн бүлэг + хэвийн болгосон оюутны код, "
          "зөвхөн энэ тестийн хосын дотор. Имэйл тааруулалтад оролцохгүй.",
          size=8.5, color="6B7793", align=WD_ALIGN_PARAGRAPH.LEFT)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
