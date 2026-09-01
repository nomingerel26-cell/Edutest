# -*- coding: utf-8 -*-
"""
Шинэ боломжуудын тест: UI маршрутууд, QR, асуулт засварлах/эрэмбэлэх,
тест хувилах/устгах, Excel ба Word экспорт, аюулгүй байдлын шалгалт.

Түр зуурын өгөгдлийн санд ажиллана — edutest.db-д хүрэхгүй.
"""

import io
import os
import sys
import tempfile
import unittest
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import database as db  # noqa: E402
import domain  # noqa: E402
import exports  # noqa: E402


class FeatureTestCase(unittest.TestCase):
    """Нэг хичээл, нэг хос, Pre+Post тест, 3 оюутны бүрэн урсгалыг бэлдэнэ."""

    @classmethod
    def setUpClass(cls):
        cls.tmpdir = tempfile.TemporaryDirectory()
        cls.db_path = os.path.join(cls.tmpdir.name, "features.db")
        db.DB_PATH = __import__("pathlib").Path(cls.db_path)
        db.init_db(cls.db_path, drop_existing=True)

        import app as app_module
        cls.app_module = app_module
        app_module.app.config.update(TESTING=True, SECRET_KEY="test-secret")

        conn = db.connect(cls.db_path)
        now = domain.now_iso()

        cls.admin_id = db.create_user(conn, "Админ", "admin@f.mn",
                                      domain.hash_password("admin1234"), "admin", "Алба", now)
        cls.teacher_id = db.create_user(conn, "Багш", "teacher@f.mn",
                                        domain.hash_password("teach1234"), "teacher", "Салбар", now)
        cls.other_id = db.create_user(conn, "Өөр багш", "other@f.mn",
                                      domain.hash_password("other1234"), "teacher", "Салбар", now)

        cls.course_id = db.create_course(conn, cls.teacher_id, "Эмийн хими 1",
                                         "PHR201", 3, "2026 Намар", now)
        cls.group_id = db.create_group(conn, cls.course_id, "PH23A", 25, now)
        cls.pair_id = db.create_pair(conn, cls.course_id, "PHR201 намрын хос", now)

        cls.tests = {}
        for kind in ("pre", "post"):
            test_id = db.create_test(
                conn, cls.course_id, cls.pair_id, cls.group_id,
                f"PHR201 {kind}", kind, "open",
                domain.generate_share_code("PHR201", kind), now,
            )
            cls.tests[kind] = test_id
            for i in range(1, 4):
                db.create_question(
                    conn, test_id, i, f"{kind} асуулт {i}",
                    {"A": "нэг", "B": "хоёр", "C": "гурав", "D": "дөрөв"}, "A", 2,
                )

        # 3 оюутан: 2 нь хоёр тестийг өгнө (тааруулагдана),
        # 1 нь зөвхөн Оролт өгнө (таараагүй мөр үүснэ).
        plan = [
            ("Батын Болд", "B230101", ["A", "A", "B"], ["A", "A", "A"]),
            ("Доржийн Сараа", "b23 0102", ["A", "B", "B"], ["A", "A", "B"]),
            ("Лхагвын Тэмүүлэн", "B230103", ["B", "B", "B"], None),
        ]
        for name, code, pre_answers, post_answers in plan:
            norm = domain.normalize_student_code(code)
            student_id = db.create_student(conn, cls.group_id, name, code, norm, None, now)
            for kind, picks in (("pre", pre_answers), ("post", post_answers)):
                if picks is None:
                    continue
                test_id = cls.tests[kind]
                attempt_id = db.create_attempt(
                    conn, test_id, cls.pair_id, student_id,
                    domain.build_match_key(cls.group_id, norm), name, now,
                )
                questions = db.list_questions(conn, test_id)
                answers = {q["id"]: picks[i] for i, q in enumerate(questions)}
                result = domain.score_attempt(questions, answers)
                db.save_answers(conn, attempt_id, result["answers"])
                db.finish_attempt(conn, attempt_id, result["total_score"],
                                  result["max_score"], result["percent"], now)

        conn.commit()
        conn.close()

    @classmethod
    def tearDownClass(cls):
        cls.tmpdir.cleanup()

    def setUp(self):
        self.client = self.app_module.app.test_client()

    def login(self, email="teacher@f.mn", password="teach1234"):
        return self.client.post("/login", data={"email": email, "password": password},
                                follow_redirects=True)

    def body(self, response):
        return response.get_data(as_text=True)


# =====================================================================
# 1. Шинэ жагсаалтын хуудсууд
# =====================================================================
class TestNavigationPages(FeatureTestCase):
    def test_all_index_pages_render(self):
        self.login()
        for path in ("/", "/courses", "/tests", "/pairs", "/results",
                     "/analytics", "/exports"):
            with self.subTest(path=path):
                r = self.client.get(path, follow_redirects=True)
                self.assertEqual(r.status_code, 200, f"{path} нээгдсэнгүй")

    def test_tests_index_lists_both_tests(self):
        self.login()
        body = self.body(self.client.get("/tests"))
        self.assertIn("PHR201 pre", body)
        self.assertIn("PHR201 post", body)

    def test_sidebar_shows_role(self):
        self.login()
        self.assertIn("Багш", self.body(self.client.get("/")))

    def test_admin_sees_users_link_teacher_does_not(self):
        self.login("admin@f.mn", "admin1234")
        self.assertIn("/admin/users", self.body(self.client.get("/")))
        self.client.get("/logout")
        self.login()
        self.assertNotIn("/admin/users", self.body(self.client.get("/")))

    def test_analytics_redirects_to_single_pair(self):
        self.login()
        r = self.client.get("/analytics", follow_redirects=False)
        self.assertEqual(r.status_code, 302)
        self.assertIn(f"/pairs/{self.pair_id}/comparison", r.headers["Location"])

    def test_dashboard_stats_present(self):
        self.login()
        body = self.body(self.client.get("/"))
        for label in ("Нийт тест", "Pre/Post хос", "Нийт оролцогч",
                      "Өнөөдөр өгсөн", "Дундаж өсөлт"):
            self.assertIn(label, body)


# =====================================================================
# 2. QR / хуваалцах
# =====================================================================
class TestSharePage(FeatureTestCase):
    def test_share_page_shows_link_and_code(self):
        self.login()
        test_id = self.tests["pre"]
        conn = db.connect(self.db_path)
        share_code = db.get_test(conn, test_id)["share_code"]
        conn.close()

        body = self.body(self.client.get(f"/tests/{test_id}/share"))
        self.assertIn(share_code, body)
        self.assertIn(f"/t/{share_code}", body)
        self.assertIn("Оюутанд илгээх холбоос", body)

    def test_qr_svg_is_valid_svg(self):
        self.login()
        r = self.client.get(f"/tests/{self.tests['pre']}/qr.svg")
        self.assertEqual(r.status_code, 200)
        self.assertIn("image/svg+xml", r.headers["Content-Type"])
        self.assertIn(b"<svg", r.data)

    def test_qr_download_sets_attachment_header(self):
        self.login()
        r = self.client.get(f"/tests/{self.tests['pre']}/qr.svg?download=1")
        self.assertIn("attachment", r.headers["Content-Disposition"])

    def test_qr_requires_login(self):
        r = self.client.get(f"/tests/{self.tests['pre']}/qr.svg", follow_redirects=False)
        self.assertEqual(r.status_code, 302)

    def test_other_teacher_cannot_open_share(self):
        self.login("other@f.mn", "other1234")
        r = self.client.get(f"/tests/{self.tests['pre']}/share")
        self.assertEqual(r.status_code, 403)


# =====================================================================
# 3. Асуулт засварлах ба эрэмбэлэх
# =====================================================================
class TestQuestionEditing(FeatureTestCase):
    def _fresh_test(self):
        """Эрэмбэ, устгалт бусад тестэд нөлөөлөхгүйн тулд тусдаа тест үүсгэнэ."""
        conn = db.connect(self.db_path)
        test_id = db.create_test(
            conn, self.course_id, None, self.group_id, "Эрэмбийн тест", "pre",
            "draft", domain.generate_share_code("PHR201", "pre"), domain.now_iso(),
        )
        ids = [
            db.create_question(conn, test_id, i, f"Асуулт {i}",
                               {"A": "a", "B": "b", "C": "c", "D": "d"}, "A", 1)
            for i in range(1, 4)
        ]
        conn.commit()
        conn.close()
        return test_id, ids

    def test_edit_question_updates_text_and_score(self):
        self.login()
        test_id, ids = self._fresh_test()
        r = self.client.post(f"/questions/{ids[0]}/edit", data={
            "text": "Шинэчилсэн асуулт", "option_a": "А сонголт", "option_b": "Б",
            "option_c": "В", "option_d": "Г", "correct_option": "C", "score": "5",
        }, follow_redirects=True)
        self.assertEqual(r.status_code, 200)

        conn = db.connect(self.db_path)
        q = db.get_question(conn, ids[0])
        conn.close()
        self.assertEqual(q["text"], "Шинэчилсэн асуулт")
        self.assertEqual(q["correct_option"], "C")
        self.assertEqual(q["score"], 5)

    def test_edit_question_rejects_invalid_input(self):
        self.login()
        test_id, ids = self._fresh_test()
        r = self.client.post(f"/questions/{ids[0]}/edit", data={
            "text": "", "option_a": "", "option_b": "б", "option_c": "в",
            "option_d": "г", "correct_option": "Z", "score": "0",
        })
        body = self.body(r)
        self.assertIn("Асуултын текст хоосон байна.", body)
        self.assertIn("Зөв хариултыг A/B/C/D-ээс сонгоно уу.", body)

    def test_move_question_swaps_order(self):
        self.login()
        test_id, ids = self._fresh_test()
        self.client.post(f"/questions/{ids[2]}/move", data={"direction": "up"},
                         follow_redirects=True)
        conn = db.connect(self.db_path)
        order = [q["id"] for q in db.list_questions(conn, test_id)]
        conn.close()
        self.assertEqual(order, [ids[0], ids[2], ids[1]])

    def test_move_first_question_up_is_noop(self):
        self.login()
        test_id, ids = self._fresh_test()
        self.client.post(f"/questions/{ids[0]}/move", data={"direction": "up"},
                         follow_redirects=True)
        conn = db.connect(self.db_path)
        order = [q["id"] for q in db.list_questions(conn, test_id)]
        conn.close()
        self.assertEqual(order, ids)

    def test_delete_renumbers_remaining_questions(self):
        self.login()
        test_id, ids = self._fresh_test()
        self.client.post(f"/questions/{ids[0]}/delete", follow_redirects=True)
        conn = db.connect(self.db_path)
        numbers = [q["order_no"] for q in db.list_questions(conn, test_id)]
        conn.close()
        self.assertEqual(numbers, [1, 2], "устгасны дараа дугаарлалт завсаргүй байх ёстой")

    def test_other_teacher_cannot_edit_question(self):
        _test_id, ids = self._fresh_test()
        self.login("other@f.mn", "other1234")
        self.assertEqual(self.client.get(f"/questions/{ids[0]}/edit").status_code, 403)


# =====================================================================
# 4. Тест хувилах ба устгах
# =====================================================================
class TestTestActions(FeatureTestCase):
    def test_duplicate_copies_questions_as_draft_with_new_code(self):
        self.login()
        source = self.tests["pre"]
        conn = db.connect(self.db_path)
        before = {t["id"] for t in db.list_all_tests(conn)}
        source_code = db.get_test(conn, source)["share_code"]
        conn.close()

        self.client.post(f"/tests/{source}/duplicate", follow_redirects=True)

        conn = db.connect(self.db_path)
        new_ids = {t["id"] for t in db.list_all_tests(conn)} - before
        self.assertEqual(len(new_ids), 1)
        copy = db.get_test(conn, new_ids.pop())
        questions = db.list_questions(conn, copy["id"])
        conn.close()

        self.assertEqual(copy["status"], "draft", "хуулбар ноорог төлөвтэй байх ёстой")
        self.assertNotEqual(copy["share_code"], source_code, "share_code шинэ байх ёстой")
        self.assertEqual(len(questions), 3)
        self.assertEqual([q["order_no"] for q in questions], [1, 2, 3])

    def test_delete_blocked_when_attempts_exist(self):
        self.login()
        r = self.client.post(f"/tests/{self.tests['pre']}/delete", follow_redirects=True)
        self.assertIn("Хариулт ирсэн тестийг устгах боломжгүй", self.body(r))

        conn = db.connect(self.db_path)
        still_there = db.get_test(conn, self.tests["pre"])
        conn.close()
        self.assertIsNotNone(still_there, "оюутны өгөгдөлтэй тест устсангүй байх ёстой")

    def test_delete_allowed_when_no_attempts(self):
        self.login()
        conn = db.connect(self.db_path)
        test_id = db.create_test(conn, self.course_id, None, None, "Устгах тест", "pre",
                                 "draft", domain.generate_share_code("X", "pre"),
                                 domain.now_iso())
        conn.commit()
        conn.close()

        self.client.post(f"/tests/{test_id}/delete", follow_redirects=True)
        conn = db.connect(self.db_path)
        gone = db.get_test(conn, test_id)
        conn.close()
        self.assertIsNone(gone)

    def test_update_test_changes_title(self):
        self.login()
        conn = db.connect(self.db_path)
        test_id = db.create_test(conn, self.course_id, None, None, "Хуучин нэр", "post",
                                 "draft", domain.generate_share_code("Y", "post"),
                                 domain.now_iso())
        conn.commit()
        conn.close()

        self.client.post(f"/tests/{test_id}/update",
                         data={"title": "Шинэ нэр", "class_group_id": str(self.group_id)},
                         follow_redirects=True)
        conn = db.connect(self.db_path)
        updated = db.get_test(conn, test_id)
        conn.close()
        self.assertEqual(updated["title"], "Шинэ нэр")
        self.assertEqual(updated["class_group_id"], self.group_id)


# =====================================================================
# 5. Excel экспорт
# =====================================================================
class TestExcelExport(FeatureTestCase):
    def test_xlsx_downloads_with_three_sheets(self):
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/comparison.xlsx")
        self.assertEqual(r.status_code, 200)
        self.assertIn("spreadsheetml", r.headers["Content-Type"])
        self.assertIn("attachment", r.headers["Content-Disposition"])
        self.assertIn(".xlsx", r.headers["Content-Disposition"])

        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(r.data))
        self.assertEqual(wb.sheetnames,
                         ["Summary", "Student Results", "Unmatched Attempts"])

    def test_xlsx_has_freeze_panes_and_filters(self):
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/comparison.xlsx")
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(r.data))
        ws = wb["Student Results"]
        self.assertEqual(ws.freeze_panes, "A2")
        self.assertIsNotNone(ws.auto_filter.ref)

    def test_xlsx_matched_and_unmatched_rows_split_correctly(self):
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/comparison.xlsx")
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(r.data))
        # 2 оюутан хоёуланг өгсөн, 1 нь зөвхөн Оролт өгсөн.
        self.assertEqual(wb["Student Results"].max_row, 3, "толгой + 2 тааруулсан мөр")
        self.assertEqual(wb["Unmatched Attempts"].max_row, 2, "толгой + 1 таараагүй мөр")

    def test_xlsx_requires_login_and_ownership(self):
        r = self.client.get(f"/pairs/{self.pair_id}/comparison.xlsx", follow_redirects=False)
        self.assertEqual(r.status_code, 302)
        self.login("other@f.mn", "other1234")
        self.assertEqual(
            self.client.get(f"/pairs/{self.pair_id}/comparison.xlsx").status_code, 403)


# =====================================================================
# 6. Word экспорт
# =====================================================================
class TestWordExport(FeatureTestCase):
    def test_docx_downloads_and_is_valid_package(self):
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/report.docx")
        self.assertEqual(r.status_code, 200)
        self.assertIn("wordprocessingml", r.headers["Content-Type"])
        self.assertIn(".docx", r.headers["Content-Disposition"])
        # .docx нь ZIP багц — уншиж чадаж байвал бүтэн.
        with zipfile.ZipFile(io.BytesIO(r.data)) as z:
            self.assertIn("word/document.xml", z.namelist())

    def test_docx_contains_summary_and_table(self):
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/report.docx")
        from docx import Document
        doc = Document(io.BytesIO(r.data))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Оролт/Гаралтын тестийн харьцуулсан тайлан", text)
        self.assertIn("Ерөнхий дүгнэлт", text)
        self.assertGreaterEqual(len(doc.tables), 3, "мета + нэгтгэл + үр дүнгийн хүснэгт")

    def test_docx_embeds_charts(self):
        """Word тайланд matplotlib-аар зурсан график шигтгэгдсэн эсэх."""
        if not exports.HAS_CHARTS:
            self.skipTest("matplotlib суулгаагүй")
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/report.docx")
        from docx import Document
        doc = Document(io.BytesIO(r.data))
        images = [rel for rel in doc.part.rels.values() if "image" in rel.reltype]
        self.assertGreaterEqual(len(images), 3, "дор хаяж 3 график байх ёстой")
        heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertIn("График", heads)

    def test_docx_still_builds_without_matplotlib(self):
        """matplotlib байхгүй ч тайлан бүрэн үүсэх ёстой (график л орохгүй)."""
        conn = db.connect(self.db_path)
        pair = db.get_pair(conn, self.pair_id)
        pre = db.list_results(conn, self.tests["pre"])
        post = db.list_results(conn, self.tests["post"])
        conn.close()
        rows = domain.match_pre_post(pre, post, self.pair_id)
        summary = domain.comparison_summary(rows)

        data = exports.build_pair_report(dict(pair), rows, summary,
                                         include_charts=False)
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Ерөнхий дүгнэлт", text)
        self.assertGreaterEqual(len(doc.tables), 3)

    def test_chart_helpers_return_none_when_no_matched_rows(self):
        """Тааруулсан оюутангүй үед график зурахгүй, алдаа ч гаргахгүй."""
        if not exports.HAS_CHARTS:
            self.skipTest("matplotlib суулгаагүй")
        only_pre = [{"status": "pre_only", "student_code": "X1",
                     "pre_percent": 50, "post_percent": None}]
        self.assertIsNone(exports.chart_students(only_pre))
        self.assertIsNone(exports.chart_distribution(only_pre))

    def test_docx_requires_ownership(self):
        self.login("other@f.mn", "other1234")
        self.assertEqual(
            self.client.get(f"/pairs/{self.pair_id}/report.docx").status_code, 403)


# =====================================================================
# 7. Файлын нэр
# =====================================================================
class TestFilenames(unittest.TestCase):
    def test_filename_pattern(self):
        name = exports.build_filename("PHR201", "Group-A", "PrePost", "xlsx",
                                      on=__import__("datetime").date(2026, 8, 31))
        self.assertEqual(name, "PHR201_Group-A_PrePost_2026-08-31.xlsx")

    def test_cyrillic_group_name_becomes_safe(self):
        name = exports.build_filename("PHR201", "ПХ23А", "PrePost", "xlsx")
        self.assertNotIn("П", name)
        self.assertTrue(name.endswith(".xlsx"))
        for bad in "/\\:*?\"<>|":
            self.assertNotIn(bad, name)

    def test_empty_group_falls_back(self):
        name = exports.build_filename("PHR201", "", "PrePost_Report", "docx")
        self.assertIn("All-Groups", name)


# =====================================================================
# 8. Аюулгүй байдал
# =====================================================================
class TestSecurity(FeatureTestCase):
    def test_student_result_hides_correct_answers(self):
        """Оюутны үр дүнгийн хуудас зөв хариултыг ил гаргах ёсгүй."""
        conn = db.connect(self.db_path)
        results = db.list_results(conn, self.tests["pre"])
        attempt_id = results[0]["attempt_id"]
        answers = db.list_answers(conn, attempt_id)
        conn.close()

        body = self.body(self.client.get(f"/r/{attempt_id}"))
        # Буруу хариулсан асуулт дээр зөв хариултын түлхүүр гарч болохгүй.
        for a in answers:
            if not a["is_correct"]:
                self.assertNotIn(f"Зөв хариулт: {a['correct_option']}", body)
        self.assertNotIn("correct_option", body)

    def test_public_student_page_blocks_draft_test(self):
        conn = db.connect(self.db_path)
        test_id = db.create_test(conn, self.course_id, None, self.group_id,
                                 "Ноорог тест", "pre", "draft",
                                 domain.generate_share_code("DRAFT", "pre"),
                                 domain.now_iso())
        code = db.get_test(conn, test_id)["share_code"]
        conn.commit()
        conn.close()
        r = self.client.get(f"/t/{code}")
        self.assertEqual(r.status_code, 403)
        self.assertIn("хараахан нээгдээгүй", self.body(r))

    def test_export_routes_reject_anonymous(self):
        for path in (f"/pairs/{self.pair_id}/comparison.csv",
                     f"/pairs/{self.pair_id}/comparison.xlsx",
                     f"/pairs/{self.pair_id}/report.docx",
                     f"/tests/{self.tests['pre']}/results.csv"):
            with self.subTest(path=path):
                r = self.client.get(path, follow_redirects=False)
                self.assertEqual(r.status_code, 302, f"{path} нэвтрэлтгүй нээгдэж байна")

    def test_download_headers_are_safe(self):
        self.login()
        r = self.client.get(f"/pairs/{self.pair_id}/comparison.xlsx")
        self.assertEqual(r.headers.get("X-Content-Type-Options"), "nosniff")
        self.assertEqual(r.headers.get("Cache-Control"), "no-store")

    def test_html_is_escaped_in_student_name(self):
        """XSS: оюутны нэр дэх HTML нь ил кодоор гарах ёсгүй."""
        conn = db.connect(self.db_path)
        payload = "<script>alert(1)</script>"
        norm = domain.normalize_student_code("B230999")
        student_id = db.create_student(conn, self.group_id, payload, "B230999",
                                       norm, None, domain.now_iso())
        attempt_id = db.create_attempt(
            conn, self.tests["pre"], self.pair_id, student_id,
            domain.build_match_key(self.group_id, norm), payload, domain.now_iso(),
        )
        db.finish_attempt(conn, attempt_id, 2, 6, 33, domain.now_iso())
        conn.commit()
        conn.close()

        self.login()
        body = self.body(self.client.get(f"/tests/{self.tests['pre']}/results"))
        self.assertNotIn("<script>alert(1)</script>", body)
        self.assertIn("&lt;script&gt;", body)

    def test_session_cookie_flags(self):
        cfg = self.app_module.app.config
        self.assertTrue(cfg["SESSION_COOKIE_HTTPONLY"])
        self.assertEqual(cfg["SESSION_COOKIE_SAMESITE"], "Lax")

    def test_no_hardcoded_secret_in_source(self):
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        with open(os.path.join(base, "app.py"), encoding="utf-8") as fh:
            source = fh.read()
        self.assertNotIn("edutest-dev-secret-key-change-me", source)

    def test_debug_is_not_enabled_by_default(self):
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        with open(os.path.join(base, "app.py"), encoding="utf-8") as fh:
            source = fh.read()
        self.assertNotIn("debug=True", source)


# =====================================================================
# 9. Шинжилгээний шүүлтүүр — тааруулалтад нөлөөлөхгүй
# =====================================================================
class TestAnalyticsFilters(FeatureTestCase):
    def test_group_filter_narrows_rows_only(self):
        self.login()
        base = self.body(self.client.get(f"/pairs/{self.pair_id}/comparison"))
        self.assertIn("B230101", base)

        filtered = self.body(self.client.get(
            f"/pairs/{self.pair_id}/comparison?status=matched"))
        # Зөвхөн Оролт өгсөн оюутан шүүгдэж хасагдана.
        self.assertIn("B230101", filtered)
        self.assertNotIn("B230103", filtered)

    def test_unknown_group_filter_shows_empty_state(self):
        self.login()
        body = self.body(self.client.get(
            f"/pairs/{self.pair_id}/comparison?group=БАЙХГҮЙ"))
        self.assertIn("Шүүлтүүрт тохирох мөр олдсонгүй", body)

    def test_matching_unaffected_by_filter(self):
        """Шүүлтүүр нь domain.match_pre_post-ийн үр дүнг өөрчлөхгүй."""
        conn = db.connect(self.db_path)
        pre = db.list_results(conn, self.tests["pre"])
        post = db.list_results(conn, self.tests["post"])
        conn.close()
        rows = domain.match_pre_post(pre, post, self.pair_id)
        matched = [r for r in rows if r["status"] == "matched"]
        self.assertEqual(len(matched), 2)
        # Хэвийн болгосон код ('b23 0102' -> 'B230102') зөв тааруулагдсан эсэх.
        self.assertIn("B230102", [r["student_code"].replace(" ", "").upper()
                                  for r in matched])


if __name__ == "__main__":
    unittest.main(verbosity=2)
